"""
Word(.docx) 본문 추출 및 단락 단위 마스킹 후 새 docx 바이트 생성.

표(table) 셀은 일반 단락과 분리되어 마스킹되기 때문에
"성명 | 홍길동" 처럼 컨텍스트 단어(헤더)가 같은 텍스트 안에 들어오지 않는다.
이 모듈은 표 처리 시 두 가지 보정을 적용한다.

1) 표 셀은 score_threshold 를 0.5 로 낮춰 KOREAN_NAME 기본 점수(0.65)가
   임계치를 통과할 수 있도록 한다.
2) 표 첫 행(헤더)에서 컨텍스트 단어가 발견되면, 같은 컬럼의 데이터 셀에
   "성명: " 같은 prefix 를 임시로 붙여 마스킹한 뒤 prefix 만 제거한다.
   Presidio 의 context boost 가 동작하여 탐지 신뢰도가 추가로 올라간다.
"""

from __future__ import annotations

import io
from typing import Iterator

import filter_engine
from docx import Document
from docx.text.paragraph import Paragraph


_BODY_THRESHOLD = 0.75
_TABLE_CELL_THRESHOLD = 0.5

_HEADER_CONTEXT_WORDS: tuple[str, ...] = (
    "성명", "이름", "담당자", "작성자", "대상자",
    "팀장", "부장", "과장", "대리", "사원",
    "이사", "본부장", "대표", "직원", "참석자",
    "인사", "직책", "직급", "역할",
)


def _iter_body_paragraphs(doc: Document) -> Iterator[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def docx_bytes_plain_text(data: bytes) -> str:
    """마스킹 엔진용으로 본문만 줄바꿈으로 이어 붙인다."""
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in _iter_body_paragraphs(doc):
        t = p.text.strip()
        if t:
            parts.append(p.text)
    return "\n".join(parts)


def _detect_header_context(text: str) -> str:
    """헤더 셀에서 컨텍스트 단어가 발견되면 'XXX: ' 형태의 prefix 를 반환한다."""
    norm = (text or "").strip()
    if not norm:
        return ""
    for w in _HEADER_CONTEXT_WORDS:
        if w in norm:
            return f"{w}: "
    return ""


def _mask_paragraph(
    p: Paragraph,
    doc_level: str,
    user_level: str,
    threshold: float,
    prefix: str = "",
) -> None:
    """단락 하나를 마스킹한다. prefix 는 컨텍스트 부스트용 임시 머리말."""
    original = p.text
    if not original or not original.strip():
        return

    target = (prefix + original) if prefix else original
    fr = filter_engine.mask_text(
        text=target,
        doc_level=doc_level,
        user_level=user_level,
        mask_open_docs=False,
        score_threshold=threshold,
    )
    masked = fr.get("masked_text") or target
    if prefix and masked.startswith(prefix):
        masked = masked[len(prefix):]

    if masked != original:
        p.text = masked


def build_masked_docx_bytes(data: bytes, doc_level: str, user_level: str) -> bytes:
    """각 단락에 filter_engine.mask_text 를 적용한 새 .docx 바이트.

    일반 단락은 기본 임계치(0.75), 표 셀은 0.5 로 낮춰 적용하며
    헤더 행 컨텍스트 단어가 있는 컬럼에는 prefix trick 으로 추가 부스트를 준다.
    """
    doc = Document(io.BytesIO(data))

    for p in doc.paragraphs:
        _mask_paragraph(p, doc_level, user_level, threshold=_BODY_THRESHOLD)

    for table in doc.tables:
        rows = list(table.rows)
        if not rows:
            continue

        header_cells = list(rows[0].cells)
        column_contexts: list[str] = [
            _detect_header_context(c.text) for c in header_cells
        ]

        for cell in header_cells:
            for p in cell.paragraphs:
                _mask_paragraph(
                    p, doc_level, user_level, threshold=_TABLE_CELL_THRESHOLD
                )

        seen_para_ids: set[int] = set()
        for row in rows[1:]:
            cells = list(row.cells)
            for idx, cell in enumerate(cells):
                ctx = column_contexts[idx] if idx < len(column_contexts) else ""
                for p in cell.paragraphs:
                    pid = id(p._p)
                    if pid in seen_para_ids:
                        continue
                    seen_para_ids.add(pid)
                    _mask_paragraph(
                        p,
                        doc_level,
                        user_level,
                        threshold=_TABLE_CELL_THRESHOLD,
                        prefix=ctx,
                    )

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
